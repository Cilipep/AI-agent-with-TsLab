from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Title
title = doc.add_heading('GridBot Final — Оптимизированная Grid-стратегия для BTC', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 1. Overview
doc.add_heading('1. Обзор стратегии', level=1)
doc.add_paragraph(
    'GridBot Final — это двусторонняя сеточная стратегия для торговли BTC фьючерсами '
    'на Binance. Стратегия использует ATR-адаптивную сетку для определения уровней '
    'входа и выхода, с реинвестированием 10% прибыли.'
)

# 2. Parameters
doc.add_heading('2. Параметры стратегии', level=1)
table = doc.add_table(rows=9, cols=3)
table.style = 'Table Grid'
headers = ['Параметр', 'Значение', 'Описание']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
data = [
    ['Инструмент', 'BTC/USDT Perpetual', 'Bitcoin фьючерсы на Binance'],
    ['Таймфрейм', '30m', '30-минутные свечи'],
    ['Начальный капитал', '40 USDT', 'Стартовый депозит'],
    ['Плечо', '10x', 'Кредитное плечо'],
    ['Grid Spacing', '2.5x ATR(14)', 'Расстояние между уровнями сетки'],
    ['Stop Distance', '3.0x ATR(14)', 'Расстояние стоп-лосса'],
    ['Risk per Trade', '10%', 'Риск на сделку от капитала'],
    ['Reinvest', '10%', 'Реинвестирование прибыли'],
]
for i, row in enumerate(data):
    for j, cell in enumerate(row):
        table.rows[i+1].cells[j].text = cell

# 3. Logic
doc.add_heading('3. Логика торговли', level=1)
doc.add_paragraph('Вход в лонг: цена закрытия ниже предыдущей → покупка по текущей цене')
doc.add_paragraph('Вход в шорт: цена закрытия выше предыдущей → продажа по текущей цене')
doc.add_paragraph('Тейк-профит: цена + GridSpacing (2.5x ATR)')
doc.add_paragraph('Стоп-лосс: цена - StopDistance (3.0x ATR)')

# 4. Backtest Results
doc.add_heading('4. Результаты бэктеста (1 месяц)', level=1)
results_table = doc.add_table(rows=12, cols=2)
results_table.style = 'Table Grid'
results = [
    ['Начальный капитал', '$40.00'],
    ['Конечный капитал', '$158.60'],
    ['Чистая прибыль', '+$118.60 (+296.5%)'],
    ['Реинвестировано', '$61.89'],
    ['Комиссии', '$61.89'],
    ['Всего сделок', '59 (28L / 31S)'],
    ['Win Rate', '66% (39W / 20L)'],
    ['Take Profit', '39'],
    ['Stop Loss', '19'],
    ['Profit Factor', '1.22'],
    ['Max Drawdown', '36.2%'],
    ['Sharpe Ratio', '6.17'],
]
for i, row in enumerate(results):
    results_table.rows[i].cells[0].text = row[0]
    results_table.rows[i].cells[1].text = row[1]

# 5. Leverage Comparison
doc.add_heading('5. Сравнение кредитных плечей', level=1)
lev_table = doc.add_table(rows=9, cols=5)
lev_table.style = 'Table Grid'
lev_headers = ['Плечо', 'PnL%', 'Drawdown', 'Sharpe', 'Комиссии']
for i, h in enumerate(lev_headers):
    lev_table.rows[0].cells[i].text = h
lev_data = [
    ['1x', '+20.2%', '4.1%', '5.72', '$2.52'],
    ['3x', '+68.3%', '12.0%', '5.81', '$9.45'],
    ['5x', '+126.2%', '19.4%', '5.91', '$19.48'],
    ['10x', '+296.5%', '36.2%', '6.17', '$61.89'],
    ['15x', '+436.6%', '50.4%', '6.46', '$129.30'],
    ['20x', '+452.7%', '65.1%', '6.81', '$204.26'],
    ['25x', '+321.3%', '80.4%', '7.21', '$252.22'],
    ['30x', '+126.2%', '90.7%', '7.70', '$247.79'],
]
for i, row in enumerate(lev_data):
    for j, cell in enumerate(row):
        lev_table.rows[i+1].cells[j].text = cell

# 6. Risk Management
doc.add_heading('6. Управление рисками', level=1)
doc.add_paragraph('• Максимальный размер позиции: 95% капитала как маржа')
doc.add_paragraph('• Стоп-лосс: 3.0x ATR от входа')
doc.add_paragraph('• Комиссия: 0.05% за сделку (Binance taker)')
doc.add_paragraph('• Реинвестирование: 10% прибыли добавляется к капиталу')

# 7. Files
doc.add_heading('7. Файлы проекта', level=1)
files_table = doc.add_table(rows=7, cols=2)
files_table.style = 'Table Grid'
files_table.rows[0].cells[0].text = 'Файл'
files_table.rows[0].cells[1].text = 'Описание'
files = [
    ['GridBot_Final.py', 'Основной код бота'],
    ['config.json', 'Конфигурация стратегии'],
    ['grid_bot_20x.py', 'Тест с разными плечами'],
    ['optimization_advanced.json', 'Результаты оптимизации'],
    ['grid_bot_binance_test.json', 'Результаты теста Binance'],
    ['GridBot_Final_Results.json', 'Финальные результаты'],
]
for i, row in enumerate(files):
    files_table.rows[i+1].cells[0].text = row[0]
    files_table.rows[i+1].cells[1].text = row[1]

# 8. How to Run
doc.add_heading('8. Запуск', level=1)
doc.add_paragraph('Бэктест: python GridBot_Final.py')
doc.add_paragraph('Live: pip install ccxt && python GridBot_Final.py live')

# 9. Recommendations
doc.add_heading('9. Рекомендации', level=1)
doc.add_paragraph('• Оптимальное плечо: 10-15x (баланс прибыли и рисков)')
doc.add_paragraph('• 20x плечо: агрессивно, DD 65%')
doc.add_paragraph('• Не превышать 20x — комиссии съедают прибыль')
doc.add_paragraph('• Мониторить drawdown, останавливать при DD > 50%')

doc.save('C:/Users/i59400f/Desktop/ai-agent/TsLabWorkspace/GridADAUSDT/GridBot_Final_Description.docx')
print('Docx создан!')
