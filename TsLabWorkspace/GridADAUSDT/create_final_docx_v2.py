from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

title = doc.add_heading('GridBot Final — Оптимизированная Grid-стратегия для BTC', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('1. Обзор стратегии', level=1)
doc.add_paragraph(
    'GridBot Final — двусторонняя сеточная стратегия для торговли BTC фьючерсами на Binance. '
    'Стратегия использует ATR-адаптивную сетку с реинвестированием 30% прибыли и трейлингом '
    'по логике Bybit (Trailing Up/Down).'
)

doc.add_heading('2. Параметры стратегии', level=1)
table = doc.add_table(rows=13, cols=3)
table.style = 'Table Grid'
headers = ['Параметр', 'Значение', 'Описание']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
data = [
    ['Символ', 'BTC/USDT:USDT', 'Bitcoin фьючерсы'],
    ['Таймфрейм', '30m', '30-минутные свечи'],
    ['Капитал', '$40 USDT', 'Начальный депозит'],
    ['Плечо', '20x', 'Кредитное плечо'],
    ['Реинвест', '30%', 'Процент реинвестирования'],
    ['Grid Spacing', '2.5x ATR(14)', 'Расстояние grid'],
    ['Stop Distance', '3.0x ATR(14)', 'Расстояние стопа'],
    ['Risk/Trade', '10%', 'Риск на сделку'],
    ['Stop Loss', '3.0x ATR', 'Стоп-лосс'],
    ['Take Profit', '2.5x ATR', 'Тейк-профит'],
    ['Trailing Up', '$70,000', 'Сверху'],
    ['Trailing Down', '$50,000', 'Снизу'],
]
for i, row in enumerate(data):
    for j, cell in enumerate(row):
        table.rows[i+1].cells[j].text = cell

doc.add_heading('3. Trailing Up/Down (Bybit)', level=1)
doc.add_paragraph('Trailing Up: когда цена достигает верхней границы сетки, сетка сдвигается вверх на 1 уровень.')
doc.add_paragraph('Trailing Down: когда цена достигает нижней границы сетки, сетка сдвигается вниз на 1 уровень.')
doc.add_paragraph('Сетка не сдвигается за пределы стоп-цен трейлинга ($70,000 и $50,000).')

doc.add_heading('4. Результаты бэктеста', level=1)
results_table = doc.add_table(rows=13, cols=2)
results_table.style = 'Table Grid'
results_table.rows[0].cells[0].text = 'Метрика'
results_table.rows[0].cells[1].text = 'Значение'
results = [
    ['Начальный капитал', '$40.00'],
    ['Конечный капитал', '$1,601.75'],
    ['Прибыль', '+$1,561.75 (+3,904%)'],
    ['Реинвестировано', '$1,723.28'],
    ['Комиссии', '$961.44'],
    ['Всего сделок', '59 (27L / 32S)'],
    ['Win Rate', '68%'],
    ['Profit Factor', '0.97'],
    ['Max Drawdown', '59.4%'],
    ['Sharpe Ratio', '9.53'],
    ['Trailing Shifts', '41↑ 40↓'],
    ['Период', '1 месяц (июнь-июль 2026)'],
]
for i, row in enumerate(results):
    results_table.rows[i].cells[0].text = row[0]
    results_table.rows[i].cells[1].text = row[1]

doc.add_heading('5. Сравнение конфигураций', level=1)
comp_table = doc.add_table(rows=6, cols=5)
comp_table.style = 'Table Grid'
comp_headers = ['Плечо', 'Реинвест', 'PnL', 'PF', 'Sharpe']
for i, h in enumerate(comp_headers):
    comp_table.rows[0].cells[i].text = h
comp_data = [
    ['10x', '10%', '+$173 (+434%)', '1.27', '7.06'],
    ['10x', '30%', '+$409 (+1022%)', '1.18', '9.02'],
    ['20x', '30%', '+$1,562 (+3,904%)', '0.97', '9.53'],
    ['20x', '45%', '+$4,278 (+10,695%)', '0.91', '10.77'],
    ['50x', '45%', '-$40 (ликвидация)', '0.68', '-'],
]
for i, row in enumerate(comp_data):
    for j, cell in enumerate(row):
        comp_table.rows[i+1].cells[j].text = cell

doc.add_heading('6. Управление рисками', level=1)
doc.add_paragraph('• Стоп-лосс: 3.0x ATR от входа')
doc.add_paragraph('• Тейк-профит: 2.5x ATR от входа')
doc.add_paragraph('• Максимальный размер позиции: 95% капитала как маржа')
doc.add_paragraph('• Комиссия: 0.05% за сделку (Binance taker)')
doc.add_paragraph('• Реинвестирование: 30% прибыли')

doc.add_heading('7. Файлы проекта', level=1)
files_table = doc.add_table(rows=7, cols=2)
files_table.style = 'Table Grid'
files_table.rows[0].cells[0].text = 'Файл'
files_table.rows[0].cells[1].text = 'Описание'
files = [
    ['GridBot_Final.py', 'Основной код бота'],
    ['config.json', 'Конфигурация'],
    ['FINAL_CONFIG.md', 'Документация параметров'],
    ['README.md', 'Обзор проекта'],
    ['GridBot_EquityCurve.png', 'График equity curve'],
    ['GridBot_TradeAnalysis.png', 'Анализ сделок'],
]
for i, row in enumerate(files):
    files_table.rows[i+1].cells[0].text = row[0]
    files_table.rows[i+1].cells[1].text = row[1]

doc.add_heading('8. Запуск', level=1)
doc.add_paragraph('Бэктест: python GridBot_Final.py')
doc.add_paragraph('Live: pip install ccxt && python GridBot_Final.py live')

doc.add_heading('9. Рекомендации', level=1)
doc.add_paragraph('• Оптимальное плечо: 10-20x (баланс прибыли и рисков)')
doc.add_paragraph('• 20x плечо + 30% реинвест = лучший баланс')
doc.add_paragraph('• 50x плечо слишком агрессивно — ликвидация')
doc.add_paragraph('• Мониторить drawdown, останавливать при DD > 50%')

doc.save('C:/Users/i59400f/Desktop/ai-agent/TsLabWorkspace/GridADAUSDT/GridBot_Final_Description.docx')
print('Docx создан!')
