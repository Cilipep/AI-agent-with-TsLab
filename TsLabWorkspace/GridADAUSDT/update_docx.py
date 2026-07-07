from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('Grid ADAUSDT - Стратегия/Grid Trading', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Section 1: Overview
doc.add_heading('1. Обзор стратегии', level=1)
doc.add_paragraph(
    'Grid ADAUSDT — это двусторонняя сеточная стратегия для торговли фьючерсами ADAUSDT '
    'на бирже Binance. Стратегия использует симметричную сетку ордеров для извлечения '
    'прибыли от колебаний цены в диапазоне.'
)

# Section 2: Parameters
doc.add_heading('2. Параметры стратегии', level=1)

table = doc.add_table(rows=10, cols=3)
table.style = 'Table Grid'

headers = ['Параметр', 'Значение', 'Описание']
for i, header in enumerate(headers):
    table.rows[0].cells[i].text = header

data = [
    ['Инструмент', 'ADAUSDT Perpetual', 'Кардано фьючерсы'],
    ['Таймфрейм', '30m', '30-минутные свечи'],
    ['Направление', 'Лонг + Шорт', 'Двусторонняя торговля'],
    ['Реинвестирование', '10%', 'Реинвестирование прибыли'],
    ['Размер позиции', '10 акций', 'Базовый размер + реинвест'],
    ['GridSpacing', 'ATR(14)', 'Адаптивная сетка на основе волатильности'],
    ['GridLevels', '3 уровня', 'Уровни сетки в каждую сторону'],
    ['StopLoss', 'ATR x 2', 'Стоп-лосс на основе волатильности'],
    ['Комиссия', '0.05%', 'Binance futures taker fee'],
]

for i, row_data in enumerate(data):
    for j, cell_data in enumerate(row_data):
        table.rows[i+1].cells[j].text = cell_data

# Section 3: Logic
doc.add_heading('3. Логика торговли', level=1)

doc.add_heading('Вход в лонг:', level=2)
doc.add_paragraph('• Цена опускается ниже текущего Close')
doc.add_paragraph('• Условие: Close < Close[1]')

doc.add_heading('Вход в шорт:', level=2)
doc.add_paragraph('• Цена поднимается выше текущего Close')
doc.add_paragraph('• Условие: Close > Close[1]')

doc.add_heading('Выход:', level=2)
doc.add_paragraph('• Тейк-профит: закрытие при достижении следующего уровня')
doc.add_paragraph('• Стоп-лосс: закрытие при пробое уровня')

# Section 4: Grid Logic
doc.add_heading('4. Grid-логика', level=1)
doc.add_paragraph(
    'Стратегия использует ATR(14) для расчёта уровней сетки. '
    'Расстояние между уровнями = ATR x GridMultiplier. '
    'Это позволяет адаптироваться к текущей волатильности рынка.'
)

# Section 5: Reinvestment
doc.add_heading('5. Реинвестирование', level=1)
doc.add_paragraph(
    'Стратегия поддерживает реинвестирование 10% прибыли. При закрытии сделки с прибылью, '
    '10% от полученного профита добавляется к размеру следующей позиции. Это позволяет '
    'увеличивать размер позиций по мере роста капитала.'
)

# Section 6: Optimization
doc.add_heading('6. Оптимизация', level=1)
doc.add_paragraph('Доступные параметры для оптимизации:')
doc.add_paragraph('• Shares (Количество): 5-20 акций')
doc.add_paragraph('• GridMultiplier: 0.5-2.0')
doc.add_paragraph('• StopMultiplier: 1.5-3.0')

# Section 7: Risk
doc.add_heading('7. Управление рисками', level=1)
doc.add_paragraph('• Стоп-лосс на каждом уровне')
doc.add_paragraph('• Максимальный размер позиции ограничен')
doc.add_paragraph('• Комиссия 0.05% (Binance futures)')

# Section 8: Files
doc.add_heading('8. Файлы стратегии', level=1)

files_table = doc.add_table(rows=10, cols=2)
files_table.style = 'Table Grid'

files_headers = ['Файл', 'Описание']
for i, header in enumerate(files_headers):
    files_table.rows[0].cells[i].text = header

files_data = [
    ['GridADAUSDT', 'Скрипт TSLab'],
    ['GridADAUSDT_Description.docx', 'Описание стратегии'],
    ['initial-strategy-description.md', 'Исходная идея'],
    ['research-hypothesis.md', 'Торговая гипотеза'],
    ['design-specification.md', 'Техническая спецификация'],
    ['execution-plan.md', 'План реализации'],
    ['risk-contract.md', 'Контракт рисков'],
    ['final-strategy-description.md', 'Полная документация'],
    ['README.md', 'Обзор проекта'],
]

for i, row_data in enumerate(files_data):
    for j, cell_data in enumerate(row_data):
        files_table.rows[i+1].cells[j].text = cell_data

# Section 9: Status
doc.add_heading('9. Статус', level=1)
doc.add_paragraph('✓ Стратегия создана в TSLab')
doc.add_paragraph('✓ Реинвестирование 10% добавлено')
doc.add_paragraph('✓ Grid-логика с ATR добавлена')
doc.add_paragraph('✓ Параметры оптимизации настроены')
doc.add_paragraph('✓ Скрипт скомпилирован и загружен')
doc.add_paragraph('⏳ Ожидание тестирования и оптимизации')

# Save
doc.save('C:/Users/i59400f/Desktop/ai-agent/TsLabWorkspace/GridADAUSDT/GridADAUSDT_Description.docx')
print('Document updated successfully!')
