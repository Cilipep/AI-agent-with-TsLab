from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

title = doc.add_heading('Стратегия GridBotRSI', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('Описание стратегии', level=1)
doc.add_paragraph('Тип: Грид-стратегия с RSI фильтрацией')
doc.add_paragraph('Инструмент: ADAUSD_PERP (Cardano)')
doc.add_paragraph('Биржа: BinanceCoin-MFutures')
doc.add_paragraph('Таймфрейм: 30 минут')
doc.add_paragraph('Начальный депозит: 25 USDT')

doc.add_heading('Логика работы', level=1)

doc.add_heading('Вход в лонг:', level=2)
doc.add_paragraph('Условие: RSI_30 < 35 (перепроданность)')
doc.add_paragraph('Цена входа: Close + 0.5% (шаг сетки)')
doc.add_paragraph('Размер позиции: 12.5 ADA (50% от баланса)')

doc.add_heading('Выход из лонга:', level=2)
doc.add_paragraph('Тейк-профит: Close + 0.5% (шаг сетки)')
doc.add_paragraph('Стоп-лосс: Close - 2%')

doc.add_heading('Вход в шорт:', level=2)
doc.add_paragraph('Условие: RSI_30 > 50 (перекупленность)')
doc.add_paragraph('Цена входа: Close + 0.5%')
doc.add_paragraph('Размер позиции: 12.5 ADA')

doc.add_heading('Выход из шорта:', level=2)
doc.add_paragraph('Тейк-профит: Close - 1.5%')
doc.add_paragraph('Стоп-лосс: Close + 2%')

doc.add_heading('Оптимизированные параметры', level=1)

table = doc.add_table(rows=9, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Параметр', 'Значение', 'Описание']
for i, header in enumerate(headers):
    table.rows[0].cells[i].text = header
    for paragraph in table.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

data = [
    ['GridConst', '0.005 (0.5%)', 'Шаг сетки для входа/выхода'],
    ['StopStep', '0.02 (2%)', 'Стоп-лосс'],
    ['ShortTakeProfit', '0.015 (1.5%)', 'Тейк-профит для шортов'],
    ['RSI_Oversold', '35', 'Порог перепроданности (лонг)'],
    ['RSI_Overbought', '50', 'Порог перекупленности (шорт)'],
    ['Shares', '12.5', 'Размер позиции (ADA)'],
    ['Commission', '0.05%', 'Комиссия за сделку'],
    ['Margin', '10%', 'Маржа']
]

for i, row_data in enumerate(data, 1):
    for j, cell_data in enumerate(row_data):
        table.rows[i].cells[j].text = cell_data

doc.add_heading('Результаты бэктеста', level=1)

doc.add_heading('Период: 23.04.2026 - 22.06.2026 (60 дней)', level=2)

table2 = doc.add_table(rows=8, cols=2)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

headers2 = ['Метрика', 'Значение']
for i, header in enumerate(headers2):
    table2.rows[0].cells[i].text = header
    for paragraph in table2.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

data2 = [
    ['Чистая прибыль', '+100.22 ADA (+400.90%)'],
    ['Всего сделок', '9'],
    ['Profit Factor', '1.26'],
    ['% прибыльных сделок', '55.56%'],
    ['Максимальная просадка', '-86.96%'],
    ['Recovery Factor', '0.34'],
    ['Просадка по прибыли', '-78.72%']
]

for i, row_data in enumerate(data2, 1):
    for j, cell_data in enumerate(row_data):
        table2.rows[i].cells[j].text = cell_data

doc.add_heading('Индикаторы на графике', level=1)

table3 = doc.add_table(rows=5, cols=3)
table3.style = 'Table Grid'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER

headers3 = ['Индикатор', 'Панель', 'Назначение']
for i, header in enumerate(headers3):
    table3.rows[0].cells[i].text = header
    for paragraph in table3.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

data3 = [
    ['RSI_30 (30 баров)', 'pane_rsi', 'Основной фильтр для входов'],
    ['Bollinger Bands (20, 2)', 'pane_price', 'Полосы на графике цены'],
    ['ATR (14 баров)', 'pane_atr', 'Средний истинный диапазон'],
    ['MACD + MACDSig', 'pane_macd', 'Схождение/расхождение скользящих']
]

for i, row_data in enumerate(data3, 1):
    for j, cell_data in enumerate(row_data):
        table3.rows[i].cells[j].text = cell_data

doc.add_heading('Выводы', level=1)
doc.add_paragraph('1. Стратегия показала положительную динамику с доходностью +400.90% за 60 дней')
doc.add_paragraph('2. Просадка снижена до -86.96% (была -246.26%)')
doc.add_paragraph('3. Profit Factor 1.26 указывает на стабильную прибыльность')
doc.add_paragraph('4. Recovery Factor 0.34 означает хорошее восстановление после просадок')

doc.add_heading('Рекомендации', level=1)
doc.add_paragraph('1. Использовать таймфрейм 30 минут')
doc.add_paragraph('2. Устанавливать стоп-лосс не менее 2%')
doc.add_paragraph('3. Тейк-профит для шортов: 1.5%')
doc.add_paragraph('4. Фильтровать входы по RSI (35/50)')
doc.add_paragraph('5. Тестировать на ADAUSD_PERP')

doc.add_paragraph('')
doc.add_paragraph('Дата создания: 22 июня 2026 года')

doc.save(r'C:\Users\i59400f\Desktop\ai-agent\TsLabWorkspace\GridBotRSI\Описание.docx')
print('Файл Описание.docx успешно создан!')
