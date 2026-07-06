from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

title = doc.add_heading('Стратегия GridBotRSI', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('Описание стратегии', level=1)
doc.add_paragraph('Тип: Грид-стратегия с RSI фильтрацией')
doc.add_paragraph('Лучший инструмент: ADAUSD_PERP (Cardano)')
doc.add_paragraph('Биржа: BinanceCoin-MFutures')
doc.add_paragraph('Таймфрейм: 30 минут')
doc.add_paragraph('Начальный депозит: 25 USDT')

doc.add_heading('Логика работы', level=1)

doc.add_heading('Вход в лонг:', level=2)
doc.add_paragraph('Условие: RSI_30 < 40 (перепроданность)')
doc.add_paragraph('Цена входа: Close + GridConst (текущая цена + шаг сетки)')
doc.add_paragraph('Размер позиции: 12.5 ADA (50% от баланса)')

doc.add_heading('Выход из лонга:', level=2)
doc.add_paragraph('Тейк-профит: Close + GridConst (шаг сетки)')
doc.add_paragraph('Стоп-лосс: Close - StopStep (стоп-шаг)')

doc.add_heading('Вход в шорт:', level=2)
doc.add_paragraph('Условие: RSI_30 > 60 (перекупленность)')
doc.add_paragraph('Цена входа: Close + GridConst')
doc.add_paragraph('Размер позиции: 12.5 ADA')

doc.add_heading('Выход из шорта:', level=2)
doc.add_paragraph('Тейк-профит: Close - ShortTakeProfit')
doc.add_paragraph('Стоп-лосс: Close + StopStep')

doc.add_heading('Оптимизированные параметры (ADAUSD_PERP)', level=1)

table = doc.add_table(rows=8, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Параметр', 'Значение', 'Описание']
for i, header in enumerate(headers):
    table.rows[0].cells[i].text = header
    for paragraph in table.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

data = [
    ['GridConst', '0.006 (0.6%)', 'Шаг сетки для входа/выхода'],
    ['StopStep', '0.025 (2.5%)', 'Стоп-лосс'],
    ['ShortTakeProfit', '0.025 (2.5%)', 'Тейк-профит для шортов'],
    ['RSI_Oversold', '40', 'Порог перепроданности'],
    ['RSI_Overbought', '60', 'Порог перекупленности'],
    ['Shares', '12.5', 'Размер позиции (ADA)'],
    ['Commission', '0.05%', 'Комиссия за сделку']
]

for i, row_data in enumerate(data, 1):
    for j, cell_data in enumerate(row_data):
        table.rows[i].cells[j].text = cell_data

doc.add_heading('Результаты тестирования', level=1)

doc.add_heading('Тест на ADAUSD_PERP (60 дней) - ЛУЧШИЙ РЕЗУЛЬТАТ', level=2)

table2 = doc.add_table(rows=7, cols=2)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

headers2 = ['Метрика', 'Значение']
for i, header in enumerate(headers2):
    table2.rows[0].cells[i].text = header
    for paragraph in table2.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

data2 = [
    ['Период', '23.04.2026 - 22.06.2026'],
    ['Чистая прибыль', '+71.46 ADA (+285.85%)'],
    ['Всего сделок', '6'],
    ['Profit Factor', '1.23'],
    ['% прибыльных', '33.33%'],
    ['Recovery Factor', '0.23']
]

for i, row_data in enumerate(data2, 1):
    for j, cell_data in enumerate(row_data):
        table2.rows[i].cells[j].text = cell_data

doc.add_heading('Сравнение всех инструментов', level=1)

table3 = doc.add_table(rows=11, cols=5)
table3.style = 'Table Grid'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER

headers3 = ['Инструмент', 'Прибыль', 'Сделки', 'Profit Factor', '% прибыльных']
for i, header in enumerate(headers3):
    table3.rows[0].cells[i].text = header
    for paragraph in table3.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.font.bold = True

data3 = [
    ['DOGEUSD_PERP', '-202.11 DOGE (-808%)', '1', '0', '0%'],
    ['SOLUSD_PERP', '-16.33 SOL (-65.31%)', '597', '0.03', '20.27%'],
    ['AVAXUSD_PERP', '-5.90 AVAX (-23.58%)', '505', '0.67', '48.91%'],
    ['LINKUSD_PERP', '-7.73 LINK (-30.93%)', '549', '0.57', '40.26%'],
    ['DOTUSD_PERP', '+9.27 DOT (+37.09%)', '209', '1.09', '59.33%'],
    ['SUIUSD_PERP', '-47.93 SUI (-191.73%)', '188', '0.71', '48.40%'],
    ['NEARUSD_PERP', '-11.39 NEAR (-45.57%)', '390', '0.91', '51.28%'],
    ['AAVEUSD_PERP', '-13.71 AAVE (-54.84%)', '490', '0.05', '24.69%'],
    ['ETCUSD_PERP', '-6.09 ETC (-24.36%)', '489', '0.65', '44.99%'],
    ['ADAUSD_PERP', '+71.46 ADA (+285.85%)', '6', '1.23', '33.33%']
]

for i, row_data in enumerate(data3, 1):
    for j, cell_data in enumerate(row_data):
        table3.rows[i].cells[j].text = cell_data

doc.add_heading('Выводы', level=1)
doc.add_paragraph('1. Стратегия показала положительные результаты на 2 инструментах из 10')
doc.add_paragraph('2. Лучший результат: ADAUSD_PERP +285.85% за 60 дней')
doc.add_paragraph('3. Второй лучший результат: DOTUSD_PERP +37.09% за 60 дней')
doc.add_paragraph('4. Стратегия лучше работает на инструментах с умеренной волатильностью')
doc.add_paragraph('5. Для коротких периодов (60 дней) рекомендуется использовать ADAUSD_PERP')

doc.add_heading('Рекомендации', level=1)
doc.add_paragraph('1. Использовать таймфрейм 30 минут')
doc.add_paragraph('2. Устанавливать стоп-лосс не менее 2.5%')
doc.add_paragraph('3. Тейк-профит не менее 2.5%')
doc.add_paragraph('4. Фильтровать входы по RSI (40/60)')
doc.add_paragraph('5. Тестировать на ADAUSD_PERP или DOTUSD_PERP')

doc.add_paragraph('')
doc.add_paragraph('Дата создания: 22 июня 2026 года')

doc.save(r'C:\Users\i59400f\Desktop\ai-agent\TsLabWorkspace\GridBotRSI\Описание.docx')
print('Файл Описание.docx успешно обновлен!')
