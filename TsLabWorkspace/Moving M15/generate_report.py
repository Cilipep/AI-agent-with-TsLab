#!/usr/bin/env python3
"""
Generate DOCX report for Moving M15 Strategy
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

def create_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('Moving M15 Strategy', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Trading Strategy Analysis Report', style='Subtitle')
    doc.add_paragraph('Generated: 2026-07-06', style='Normal')
    
    # Section 1: Strategy Overview
    doc.add_heading('1. Strategy Overview', level=1)
    
    doc.add_paragraph(
        'The Moving M15 Strategy is a trend-following approach combining SMMA crossover '
        'with Williams %R oscillator for momentum confirmation. The strategy was analyzed '
        'across multiple cryptocurrency instruments and timeframes to identify optimal '
        'configurations.'
    )
    
    doc.add_heading('Strategy Rules', level=2)
    doc.add_paragraph('Main Window:', style='List Bullet')
    doc.add_paragraph('SMMA(X) on Close vs SMMA(X) shifted X bars', style='List Bullet 2')
    doc.add_paragraph('Sub Window:', style='List Bullet')
    doc.add_paragraph('Williams %R(Y) oscillator', style='List Bullet 2')
    doc.add_paragraph('SMMA(8) of Williams %R', style='List Bullet 2')
    doc.add_paragraph('SMMA(21) of Williams %R', style='List Bullet 2')
    doc.add_paragraph('Entry Conditions:', style='List Bullet')
    doc.add_paragraph('BUY: SMMA crosses above shifted SMMA + WR crosses above WR8 + WR > WR21', style='List Bullet 2')
    doc.add_paragraph('SELL: SMMA crosses below shifted SMMA + WR crosses below WR8 + WR < WR21', style='List Bullet 2')
    
    # Section 2: Performance Results
    doc.add_heading('2. Performance Results', level=1)
    
    doc.add_heading('Top Performing Instruments (H1, default params)', level=2)
    
    table = doc.add_table(rows=11, cols=9)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['Rank', 'Symbol', 'SL%', 'TP%', 'SMMA', 'WR', 'Trades', 'Win Rate%', 'Return%']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    data = [
        ['1', 'SOL', '3.0', '8.0', '7', '10', '18', '50.0', '+49.84'],
        ['2', 'LINK', '3.0', '8.0', '5', '10', '37', '35.1', '+30.05'],
        ['3', 'INJ', '3.0', '6.0', '7', '10', '21', '38.1', '+10.43'],
        ['4', 'ETH', '3.0', '8.0', '7', '10', '20', '35.0', '+8.38'],
        ['5', 'DOGE', '3.0', '6.0', '7', '10', '12', '41.7', '+5.57'],
        ['6', 'FET', '3.0', '6.0', '7', '10', '27', '33.3', '+2.58'],
        ['7', 'WIF', '3.0', '6.0', '7', '10', '25', '36.0', '-0.84'],
        ['8', 'AVAX', '3.0', '6.0', '7', '10', '20', '35.0', '-2.76'],
        ['9', 'BTC', '3.0', '8.0', '7', '10', '19', '5.3', '-21.81'],
        ['10', 'BONK', '3.0', '6.0', '7', '10', '25', '28.0', '-16.94'],
    ]
    
    for row_idx, row_data in enumerate(data):
        for col_idx, cell_data in enumerate(row_data):
            table.rows[row_idx + 1].cells[col_idx].text = cell_data
    
    doc.add_paragraph()
    
    doc.add_heading('Optimized Results with EMA 200 Filter (H1, 3mo)', level=2)
    
    table2 = doc.add_table(rows=7, cols=8)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers2 = ['Symbol', 'SL%', 'TP%', 'SMMA', 'WR', 'EMA', 'Trades', 'Return%']
    for i, header in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    data2 = [
        ['DOGE', '3.0', '8.0', '5', '10', '200', '29', '+41.08'],
        ['SOL', '3.0', '6.0', '5', '10', '200', '43', '+19.73'],
        ['FET', '3.0', '4.0', '7', '10', '200', '48', '+19.48'],
        ['LINK', '2.0', '6.0', '5', '10', '200', '37', '+18.95'],
        ['BTC', '2.0', '8.0', '5', '10', '200', '19', '+10.48'],
        ['ETH', '2.0', '8.0', '7', '10', '200', '19', '+7.33'],
    ]
    
    for row_idx, row_data in enumerate(data2):
        for col_idx, cell_data in enumerate(row_data):
            table2.rows[row_idx + 1].cells[col_idx].text = cell_data
    
    doc.add_paragraph()
    
    # Section 3: Timeframe Analysis
    doc.add_heading('3. Timeframe Analysis', level=1)
    
    doc.add_paragraph(
        'The strategy was tested across multiple timeframes to identify optimal trading intervals.'
    )
    
    table3 = doc.add_table(rows=6, cols=4)
    table3.style = 'Table Grid'
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers3 = ['Timeframe', 'Profitable Instruments', 'Avg Return%', 'Best Instrument']
    for i, header in enumerate(headers3):
        cell = table3.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    data3 = [
        ['D1', '7/7 (100%)', '+20.83%', 'INJ +31.79%'],
        ['H4', '7/7 (100%)', '+12.34%', 'LINK +28.24%'],
        ['H1', '6/7 (86%)', '+12.08%', 'SOL +49.84%'],
        ['M15', '6/7 (86%)', '~20%', 'FET +51.17%'],
        ['M5', '2/5 (40%)', '-5.47%', 'ETH +18.24%'],
    ]
    
    for row_idx, row_data in enumerate(data3):
        for col_idx, cell_data in enumerate(row_data):
            table3.rows[row_idx + 1].cells[col_idx].text = cell_data
    
    doc.add_paragraph()
    
    # Section 4: Key Findings
    doc.add_heading('4. Key Findings', level=1)
    
    findings = [
        'SOL demonstrates the strongest performance with 50% win rate and 1:2.7 R:R ratio',
        'SMMA period 7 outperforms period 5 across most instruments',
        'TP=8% works better than TP=6% for higher-volatility assets',
        'EMA 200 trend filter makes strategy profitable across all tested instruments',
        'BTC and meme tokens (BONK) should be avoided - strategy fails in those conditions',
        'D1 and H4 timeframes show most consistent results (100% profitable)',
        'M5 timeframe too noisy - only 40% of instruments profitable',
    ]
    
    for finding in findings:
        doc.add_paragraph(finding, style='List Number')
    
    # Section 5: Portfolio Allocation
    doc.add_heading('5. Recommended Portfolio Allocation', level=1)
    
    doc.add_paragraph('Primary (highest confidence):', style='List Bullet')
    doc.add_paragraph('SOL - highest Sharpe, lowest MDD', style='List Bullet 2')
    
    doc.add_paragraph('Secondary:', style='List Bullet')
    doc.add_paragraph('LINK, INJ - solid performers with good risk/reward', style='List Bullet 2')
    
    doc.add_paragraph('Tertiary:', style='List Bullet')
    doc.add_paragraph('ETH, DOGE, FET - profitable but higher volatility', style='List Bullet 2')
    
    doc.add_paragraph('Avoid:', style='List Bullet')
    doc.add_paragraph('BTC, BONK, WIF, AVAX - strategy underperforms', style='List Bullet 2')
    
    # Section 6: TSLab Implementation
    doc.add_heading('6. TSLab Implementation', level=1)
    
    doc.add_paragraph(
        'A custom Williams %R indicator was created and uploaded to TSLab as a C# DLL. '
        'The strategy was implemented with the following blocks:'
    )
    
    blocks = [
        'WilliamsR (custom indicator) - Momentum oscillator',
        'SMMA(5) - Main signal line',
        'SMMA(5) shifted - Signal comparison line',
        'EMA(200) - Trend filter',
        'CrossOver/CrossUnder - Signal detection',
        'Greater/Less - Condition filters',
        'And - Signal combination',
        'OpenPositionByMarket - Entry execution',
        'ClosePositionByStopItem - Stop loss protection',
        'ClosePositionByProfitItem - Take profit target',
    ]
    
    for block in blocks:
        doc.add_paragraph(block, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('DLL File: WilliamsR.dll (uploaded to TSLab)')
    doc.add_paragraph('Script ID: 107 (Moving M15 DOGE EMA200)')
    
    # Section 7: Files
    doc.add_heading('7. Generated Files', level=1)
    
    files_table = doc.add_table(rows=8, cols=2)
    files_table.style = 'Table Grid'
    files_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    files_headers = ['File', 'Description']
    for i, header in enumerate(files_headers):
        cell = files_table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    files_data = [
        ['backtest.py', 'Core backtest engine with SMMA/WR indicators'],
        ['optimize_sol.py', 'SOL parameter optimization'],
        ['optimize_link.py', 'LINK parameter optimization'],
        ['backtest_sol.py', 'SOL detailed backtest'],
        ['portfolio.py', 'Multi-instrument portfolio analysis'],
        ['SUMMARY.md', 'Strategy summary document'],
        ['WilliamsR.cs', 'Custom TSLab indicator source code'],
    ]
    
    for row_idx, row_data in enumerate(files_data):
        for col_idx, cell_data in enumerate(row_data):
            files_table.rows[row_idx + 1].cells[col_idx].text = cell_data
    
    # Save document
    output_path = r"C:\Users\i59400f\Desktop\ai-agent\TsLabWorkspace\Moving M15\Moving_M15_Strategy_Report.docx"
    doc.save(output_path)
    print(f"Report saved to: {output_path}")

if __name__ == "__main__":
    create_report()
