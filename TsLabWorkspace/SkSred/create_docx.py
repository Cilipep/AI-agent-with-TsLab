from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

doc = Document()

# Title
title = doc.add_heading('SkSred v2.0 — Оптимизированная стратегия', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Date
doc.add_paragraph(f'Дата: {datetime.now().strftime("%d.%m.%Y")}')

# Overview
doc.add_heading('1. Обзор стратегии', level=1)
doc.add_paragraph(
    'SkSred — это автоматическая торговая стратегия на основе пересечения скользящих средних (SMA) '
    'с адаптивным реинвестированием прибыли по тренду. Стратегия предназначена для торговли '
    'криптовалютными фьючерсами на Binance.'
)

# Parameters
doc.add_heading('2. Параметры стратегии', level=1)

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Параметр'
hdr_cells[1].text = 'Значение'

params = [
    ('Инструмент', 'DOGEUSD_PERP (Binance Futures)'),
    ('Таймфрейм', '1H (часовой)'),
    ('Начальный капитал', '25 USDT'),
    ('Плечо', '1.7x'),
    ('Risk на сделку', '0.3%'),
    ('Stop Loss', '1.6x ATR'),
    ('Take Profit', '2.5x ATR'),
    ('Trailing Stop', '0.8%'),
    ('Комиссия', '0.1%'),
]

for param, value in params:
    row_cells = table.add_row().cells
    row_cells[0].text = param
    row_cells[1].text = value

# Indicators
doc.add_heading('3. Индикаторы', level=1)

table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Table Grid'
hdr_cells2 = table2.rows[0].cells
hdr_cells2[0].text = 'Индикатор'
hdr_cells2[1].text = 'Период'
hdr_cells2[2].text = 'Назначение'

indicators = [
    ('SMA Fast', '15', 'Краткосрочная скользящая средняя'),
    ('SMA Slow', '40', 'Долгосрочная скользящая средняя'),
    ('EMA Trend', '200', 'Фильтр тренда'),
    ('RSI', '14', 'Фильтр перекупленности/перепроданности'),
    ('ATR', '14', 'Расчёт стоп-лосс и тейк-профит'),
]

for ind, period, purpose in indicators:
    row_cells = table2.add_row().cells
    row_cells[0].text = ind
    row_cells[1].text = period
    row_cells[2].text = purpose

# Entry signals
doc.add_heading('4. Сигналы входа', level=1)

doc.add_heading('LONG:', level=2)
doc.add_paragraph('• SMA Fast пересекает SMA Slow снизу вверх (Golden Cross)', style='List Bullet')
doc.add_paragraph('• RSI < 70 (нет перекупленности)', style='List Bullet')
doc.add_paragraph('• Цена выше EMA200 (восходящий тренд)', style='List Bullet')

doc.add_heading('SHORT:', level=2)
doc.add_paragraph('• SMA Fast пересекает SMA Slow сверху вниз (Death Cross)', style='List Bullet')
doc.add_paragraph('• RSI > 30 (нет перепроданности)', style='List Bullet')
doc.add_paragraph('• Цена ниже EMA200 (нисходящий тренд)', style='List Bullet')

# Risk management
doc.add_heading('5. Управление рисками', level=1)

doc.add_paragraph('• Stop Loss: 1.6x ATR от точки входа', style='List Bullet')
doc.add_paragraph('• Take Profit: 2.5x ATR от точки входа', style='List Bullet')
doc.add_paragraph('• Trailing Stop: 0.8% от пика прибыли', style='List Bullet')
doc.add_paragraph('• Максимальная просадка: 25%', style='List Bullet')

# Reinvestment
doc.add_heading('6. Адаптивное реинвестирование', level=1)

table3 = doc.add_table(rows=1, cols=3)
table3.style = 'Table Grid'
hdr_cells3 = table3.rows[0].cells
hdr_cells3[0].text = 'Тренд'
hdr_cells3[1].text = 'Реинвест'
hdr_cells3[2].text = 'Условие'

reinvest = [
    ('Восходящий', '8%', 'Цена > EMA200 * 1.02'),
    ('Боковой', '2%', 'EMA200 * 0.98 ≤ Цена ≤ EMA200 * 1.02'),
    ('Нисходящий', '0%', 'Цена < EMA200 * 0.98'),
]

for trend, pct, condition in reinvest:
    row_cells = table3.add_row().cells
    row_cells[0].text = trend
    row_cells[1].text = pct
    row_cells[2].text = condition

# Backtest results
doc.add_heading('7. Результаты бэктеста', level=1)

table4 = doc.add_table(rows=1, cols=2)
table4.style = 'Table Grid'
hdr_cells4 = table4.rows[0].cells
hdr_cells4[0].text = 'Метрика'
hdr_cells4[1].text = 'Значение'

results = [
    ('Период', '2026-01-08 — 2026-07-07 (6 мес.)'),
    ('Начальный капитал', '25 USDT'),
    ('Конечный капитал', '33.13 USDT'),
    ('Чистая прибыль', '+8.13 USDT (+32.51%)'),
    ('Максимальная просадка', '23.71%'),
    ('Sharpe Ratio', '1.37'),
    ('Profit Factor', '1.49'),
    ('Win Rate', '50.0%'),
    ('Всего сделок', '38'),
    ('Прибыльных', '19'),
    ('Убыточных', '19'),
]

for metric, value in results:
    row_cells = table4.add_row().cells
    row_cells[0].text = metric
    row_cells[1].text = value

# By trends
doc.add_heading('8. Результаты по трендам', level=1)

table5 = doc.add_table(rows=1, cols=3)
table5.style = 'Table Grid'
hdr_cells5 = table5.rows[0].cells
hdr_cells5[0].text = 'Тренд'
hdr_cells5[1].text = 'Сделок'
hdr_cells5[2].text = 'PnL'

trends = [
    ('Восходящий (8% реинвест)', '11', '+$4.97'),
    ('Боковой (2% реинвест)', '6', '-$8.82'),
    ('Нисходящий (0% реинвест)', '21', '+$14.15'),
]

for trend, trades, pnl in trends:
    row_cells = table5.add_row().cells
    row_cells[0].text = trend
    row_cells[1].text = trades
    row_cells[2].text = pnl

# Files
doc.add_heading('9. Файлы проекта', level=1)

files = [
    'backtest_final_optimized.py — Оптимизированный бэктест',
    'strategy_config.json — Конфигурация стратегии',
    'SkSred.tscript — Визуальный скрипт TSLab',
    'CIRCUIT_BREAKER.md — Документация circuit breaker',
]

for f in files:
    doc.add_paragraph(f, style='List Bullet')

# Recommendations
doc.add_heading('10. Рекомендации', level=1)

doc.add_paragraph('• Начинать с малого капитала ($5-10) для тестирования', style='List Bullet')
doc.add_paragraph('• Мониторить первые 24-48 часов', style='List Bullet')
doc.add_paragraph('• Настроить алерты на сделки', style='List Bullet')
doc.add_paragraph('• Иметь стоп-лосс на уровне портфеля', style='List Bullet')
doc.add_paragraph('• Не превышать плечо 1.7x для контроля просадки', style='List Bullet')

# Save
filename = 'C:\\Users\\i59400f\\Desktop\\ai-agent\\TsLabWorkspace\\SkSred\\SkSred_v2_Description.docx'
doc.save(filename)
print(f'Документ сохранён: {filename}')
