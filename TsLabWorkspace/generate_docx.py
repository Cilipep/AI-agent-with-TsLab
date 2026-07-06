from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

doc.add_heading('Настройка автономной торговли TSLab на VPS', 0)

doc.add_heading('1. Выбор VPS', 1)
doc.add_heading('Рекомендуемые характеристики:', 2)
table = doc.add_table(rows=6, cols=2)
table.style = 'Light Grid Accent 1'
data = [
    ('OS', 'Windows Server 2019/2022 (64-bit)'),
    ('RAM', 'от 4 ГБ (8 ГБ рекомендуется)'),
    ('CPU', '2+ ядра'),
    ('Диск', '50+ ГБ SSD'),
    ('Сеть', 'стабильный интернет, низкий пинг до биржи'),
    ('', ''),
]
for i, (k, v) in enumerate(data):
    if k:
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v

doc.add_heading('Провайдеры:', 2)
for p in ['Timeweb Cloud', 'Selectel', 'DigitalOcean', 'AWS (EC2)', 'Azure (VM)']:
    doc.add_paragraph(p, style='List Bullet')

doc.add_heading('2. Установка TSLab', 1)
steps = [
    'Подключитесь к VPS через RDP',
    'Скачайте TSLab Console с официального сайта',
    'Запустите установщик',
    'Выберите тип установки: Консоль (не Десктоп)',
    'Укажите порт API (по умолчанию 5000)',
    'Завершите установку',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_heading('3. Настройка TSLab', 1)
doc.add_heading('3.1. Настройка API', 2)
doc.add_paragraph('Файл → Настройки → Web API')
doc.add_paragraph('Порт: 5000', style='List Bullet')
doc.add_paragraph('Разрешить внешние подключения: ✓', style='List Bullet')

doc.add_heading('3.2. Настройка брокера', 2)
doc.add_paragraph('Файл → Настройки → Брокеры')
doc.add_paragraph('Выберите биржу (Binance, Bybit и т.д.)', style='List Bullet')
doc.add_paragraph('Введите API ключ и Secret', style='List Bullet')
doc.add_paragraph('Нажмите "Тест подключения"', style='List Bullet')

doc.add_heading('3.3. Настройка автозапуска', 2)
doc.add_paragraph('Файл → Настройки → Общие')
doc.add_paragraph('Автозапуск при старте Windows: ✓', style='List Bullet')
doc.add_paragraph('Автозапуск агентов: ✓', style='List Bullet')

doc.add_heading('4. Перенос скрипта', 1)
doc.add_heading('4.1. Экспорт с локального ПК', 2)
code = doc.add_paragraph()
code.style = 'No Spacing'
run = code.add_run('curl.exe -s http://localhost:5000/api/scripts/GridBotRSI/json > GridBotRSI.json')
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_heading('4.2. Импорт на VPS', 2)
code = doc.add_paragraph()
code.style = 'No Spacing'
run = code.add_run('curl.exe -X POST http://localhost:5000/api/scripts -H "Content-Type: application/json" --data-binary "@GridBotRSI.json"')
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_heading('5. Запуск агента', 1)
code = doc.add_paragraph()
code.style = 'No Spacing'
run = code.add_run('curl.exe -X POST http://localhost:5000/api/agents/GridBotRSI-Live/start')
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_paragraph('Проверка статуса:')
code = doc.add_paragraph()
code.style = 'No Spacing'
run = code.add_run('curl.exe -s http://localhost:5000/api/agent-manager/agents')
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_heading('6. Мониторинг', 1)
doc.add_heading('6.1. Через API', 2)
for cmd in [
    'curl.exe -s http://localhost:5000/api/trading/own-positions',
    'curl.exe -s http://localhost:5000/api/trading/own-trades',
    'curl.exe -s http://localhost:5000/api/scripts/GridBotRSI/logs',
]:
    code = doc.add_paragraph()
    code.style = 'No Spacing'
    run = code.add_run(cmd)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

doc.add_heading('6.2. Через Telegram-бот (опционально)', 2)
for s in [
    'Создайте бота через @BotFather',
    'В TSLab: Настройки → Уведомления → Telegram',
    'Введите Token и Chat ID',
]:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('7. Безопасность', 1)
doc.add_heading('7.1. Firewall', 2)
code = doc.add_paragraph()
code.style = 'No Spacing'
run = code.add_run('New-NetFirewallRule -DisplayName "TSLab API" -Direction Inbound -LocalPort 5000 -Protocol TCP -Action Allow')
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_heading('7.2. Ограничение доступа', 2)
for s in [
    'Используйте статический IP для подключения',
    'Настройте RDP на нестандартный порт',
    'Включите 2FA для Windows',
]:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('7.3. Резервное копирование', 2)
code = doc.add_paragraph()
code.style = 'No Spacing'
run = code.add_run('curl.exe -s http://localhost:5000/api/scripts/GridBotRSI/json > backup_$(Get-Date -Format yyyyMMdd).json')
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_heading('8. Автозапуск при перезагрузке VPS', 1)
doc.add_heading('8.1. Через планировщик заданий', 2)
code = doc.add_paragraph()
code.style = 'No Spacing'
run = code.add_run('''$action = New-ScheduledTaskAction -Execute "C:\\Program Files\\TSLab\\TSLab 3.0\\TSLab.Console.exe"
$trigger = New-ScheduledTaskTrigger -AtStartup
Register-ScheduledTask -TaskName "TSLab AutoStart" -Action $action -Trigger $trigger -RunLevel Highest''')
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_heading('8.2. Через реестр', 2)
code = doc.add_paragraph()
code.style = 'No Spacing'
run = code.add_run('HKLM\\SOFTWARE\\Microsoft\\Windows\\CurrentVersion\\Run\nTSLab = "C:\\Program Files\\TSLab\\TSLab 3.0\\TSLab.Console.exe"')
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_heading('9. Рекомендации', 1)
for s in [
    'Регулярно проверяйте логи на ошибки',
    'Настройте уведомления о критических событиях',
    'Делайте бэкапы скриптов перед изменениями',
    'Мониторинг пинга до биржи',
    'Используйте статический IP для API-ключей биржи',
]:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('10. Частые проблемы', 1)
table = doc.add_table(rows=5, cols=2)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Проблема'
table.rows[0].cells[1].text = 'Решение'
problems = [
    ('Агент не запускается', 'Проверьте настройки брокера и баланс'),
    ('Нет сделок', 'Проверьте условия входа в скрипте'),
    ('Высокий пинг', 'Смените регион VPS'),
    ('Ошибки API', 'Проверьте Firewall и порты'),
]
for i, (k, v) in enumerate(problems, 1):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('Дата создания: 24.06.2026\nВерсия TSLab: 3.0')
run.font.size = Pt(9)
run.font.color.rgb = None

doc.save(r'C:\Users\i59400f\Desktop\ai-agent\TsLabWorkspace\Настройки Автономной торговли.docx')
print('Done')
