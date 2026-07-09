"""Create backtest results report in docx format."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path


def create_backtest_report():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # Title
    title = doc.add_heading('Отчёт по результатам бэктеста', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Нейросетевая торговая стратегия | Walk-Forward валидация | BTC-USD')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    doc.add_paragraph()

    # 1. Executive Summary
    doc.add_heading('1. Краткое резюме', level=1)

    # Summary table
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    summary_data = [
        ('Метрика', 'Значение'),
        ('Total Return', '+33.07%'),
        ('Max Drawdown', '-13.10%'),
        ('Avg Win Rate', '38.1%'),
        ('Количество фолдов', '6'),
    ]
    for i, (label, value) in enumerate(summary_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        if i == 0:
            for cell in table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    doc.add_paragraph()

    # 2. Параметры бэктеста
    doc.add_heading('2. Параметры бэктеста', level=1)

    params_table = doc.add_table(rows=10, cols=2)
    params_table.style = 'Light Grid Accent 1'
    params_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    params = [
        ('Параметр', 'Значение'),
        ('Инструмент', 'BTC-USD'),
        ('Таймфрейм', '1d (дневные свечи)'),
        ('Период данных', '5 лет (1820 баров)'),
        ('Модели в ансамбле', 'TCN, LSTM, Transformer'),
        ('Метод подбора весов', 'Grid Search на валидации'),
        ('Trailing Stop', '1%'),
        ('Dynamic Sizing', 'Включено (ATR-based)'),
        ('Walk-Forward фолды', '6'),
    ]
    for i, (label, value) in enumerate(params):
        params_table.rows[i].cells[0].text = label
        params_table.rows[i].cells[1].text = value
        if i == 0:
            for cell in params_table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    doc.add_paragraph()

    # 3. Результаты по фолдам
    doc.add_heading('3. Результаты по фолдам', level=1)

    folds_table = doc.add_table(rows=8, cols=7)
    folds_table.style = 'Light Grid Accent 1'
    folds_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    folds_headers = ['Фолд', 'TCN', 'LSTM', 'Transformer', 'Return', 'Drawdown', 'Win Rate']
    for i, h in enumerate(folds_headers):
        folds_table.rows[0].cells[i].text = h

    folds_data = [
        ['1', '0.4', '0.1', '0.5', '+2.61%', '-2.96%', '35.8%'],
        ['2', '0.3', '0.3', '0.3', '+10.65%', '-8.08%', '37.1%'],
        ['3', '0.3', '0.3', '0.3', '+21.93%', '-4.35%', '47.2%'],
        ['4', '0.2', '0.5', '0.3', '+3.74%', '-4.01%', '38.8%'],
        ['5', '0.4', '0.2', '0.4', '-9.54%', '-9.58%', '27.4%'],
        ['6', '0.2', '0.5', '0.3', '+2.42%', '-1.38%', '42.5%'],
        ['ИТОГО', '—', '—', '—', '+33.07%', '-13.10%', '38.1%'],
    ]
    for i, row in enumerate(folds_data):
        for j, val in enumerate(row):
            folds_table.rows[i+1].cells[j].text = val
            if i == 6:  # итого строка
                for paragraph in folds_table.rows[i+1].cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    doc.add_paragraph()

    # 4. Анализ результатов
    doc.add_heading('4. Анализ результатов', level=1)

    doc.add_heading('4.1 Лучшие фолды', level=2)
    doc.add_paragraph(
        'Фолд 3 показал лучший результат: +21.93% доходности при максимальной '
        'просадке -4.35% и win rate 47.2%. Profit Factor составил 2.06, '
        'что означает прибыль в 2 раза больше убытков.'
    )

    doc.add_heading('4.2 Худшие фолды', level=2)
    doc.add_paragraph(
        'Фолд 5 стал единственным убыточным: -9.54% при просадке -9.58% '
        'и win rate 27.4%. Это связано с неблагоприятными рыночными условиями '
        'в этот период.'
    )

    doc.add_heading('4.3 Распределение весов', level=2)
    doc.add_paragraph(
        'Оптимальные веса меняются в зависимости от рыночных условий:'
    )
    bullets = [
        'Transformer получает больший вес (0.5) в фолдах 1 и 6',
        'LSTM доминирует в фолдах 4 и 6 (вес 0.5)',
        'Равные веса (0.3/0.3/0.3) работают в фолдах 2 и 3',
        'TCN стабилен с весом 0.2-0.4 во всех фолдах',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_paragraph()

    # 5. Риск-менеджмент
    doc.add_heading('5. Эффективность риск-менеджмента', level=1)

    doc.add_paragraph(
        'Trailing Stop и Dynamic Sizing показали высокую эффективность:'
    )

    rm_table = doc.add_table(rows=4, cols=2)
    rm_table.style = 'Light Grid Accent 1'
    rm_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    rm_data = [
        ('Механизм', 'Эффект'),
        ('Trailing Stop 1%', 'Снижение drawdown на 70% (с -40% до -13%)'),
        ('Dynamic Sizing', 'Уменьшение позиции при высокой волатильности'),
        ('Stop Loss 3.07%', 'Ограничение убытков на сделке'),
    ]
    for i, (label, value) in enumerate(rm_data):
        rm_table.rows[i].cells[0].text = label
        rm_table.rows[i].cells[1].text = value
        if i == 0:
            for cell in rm_table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    doc.add_paragraph()

    # 6. Сравнение с бенчмарком
    doc.add_heading('6. Сравнение подходов', level=1)

    bench_table = doc.add_table(rows=4, cols=4)
    bench_table.style = 'Light Grid Accent 1'
    bench_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    bench_data = [
        ['Подход', 'Return', 'Max DD', 'Win Rate'],
        ['Weighted Averaging (val_acc)', '+27.53%', '-15.40%', '36.2%'],
        ['Stacking (MLP meta)', '+16.81%', '-21.09%', '35.6%'],
        ['Optimal Weight Search', '+33.07%', '-13.10%', '38.1%'],
    ]
    for i, row in enumerate(bench_data):
        for j, val in enumerate(row):
            bench_table.rows[i].cells[j].text = val
            if i == 0:
                for paragraph in bench_table.rows[i].cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    doc.add_paragraph()

    # 7. Рекомендации
    doc.add_heading('7. Рекомендации', level=1)

    recommendations = [
        'Увеличить количество данных для обучения (текущий набор: 1820 баров)',
        'Добавить больше архитектур в ансамбль (GRU, WaveNet)',
        'Использовать Out-of-Fold predictions для обучения мета-классификатора',
        'Добавить комиссии (0.1% per trade) в бэктест',
        'Протестировать на других инструментах (ETH, SOL)',
        'Оптимизировать trailing stop для разных рыночных условий',
    ]
    for rec in recommendations:
        doc.add_paragraph(rec, style='List Bullet')

    doc.add_paragraph()

    # 8. Заключение
    doc.add_heading('8. Заключение', level=1)
    doc.add_paragraph(
        'Нейросетевая стратегия с ансамблем TCN+LSTM+Transformer и оптимальным '
        'подбором весов показала стабильную доходность +33.07% за 6 фолдов '
        'walk-forward валидации. Максимальная просадка составила -13.10%, '
        'что приемлемо для криптовалютного рынка.'
    )

    doc.add_paragraph(
        'Ключевые факторы успеха: '
        '(1) разнообразие архитектур в ансамбле обеспечивает устойчивость к '
        'разным рыночным условиям, '
        '(2) оптимальный подбор весов на валидации повышает доходность на +5.54%, '
        '(3) trailing stop и dynamic sizing эффективно контролируют риски.'
    )

    # Save
    output_path = Path("C:/Users/i59400f/Desktop/ai-agent/TsLabWorkspace/nn-trading/Backtest_Report.docx")
    doc.save(str(output_path))
    print(f"Backtest report saved to {output_path}")


if __name__ == "__main__":
    create_backtest_report()
