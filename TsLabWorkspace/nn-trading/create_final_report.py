"""Create final DOCX report with all improvements."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
from datetime import datetime


def create_final_report():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # Styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading('Финальный отчёт', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Neural Network Trading Strategy - Enhanced Version')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(89, 89, 89)
    run.font.italic = True

    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f'Дата: {datetime.now().strftime("%d.%m.%Y")}')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(89, 89, 89)

    doc.add_page_break()

    # Executive Summary
    doc.add_heading('Резюме', level=1)
    doc.add_paragraph(
        'В результате комплексной оптимизации нейросетевой торговой стратегии '
        'были достигнуты значительные улучшения:'
    )

    # Key results table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['Метрика', 'Было', 'Стало']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        for paragraph in table.rows[0].cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    results = [
        ('Total Return', '-67.10%', '+26.62%'),
        ('Max Drawdown', '-67.22%', '-4.08%'),
        ('Win Rate', '31.2%', '30.7%'),
    ]

    for metric, before, after in results:
        row_cells = table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = before
        row_cells[2].text = after

    doc.add_page_break()

    # Development Stages
    doc.add_heading('Этапы разработки', level=1)

    stages = [
        ('1. Базовая версия', 'LSTM модель с 30+ индикаторами', '-67.10%', 'Переобучение'),
        ('2. Новые признаки', 'Stochastic RSI, Aroon, Keltner', '+12.30%', 'Улучшение'),
        ('3. TA-Lib интеграция', '60+ индикаторов TA-Lib', '+11.91%', 'Стабильность'),
        ('4. Trailing Stop', '1% trailing stop + dynamic sizing', '+16.36%', 'Снижение рисков'),
        ('5. TCN модель', 'Temporal Convolutional Network', '+26.62%', 'Лучшая модель'),
    ]

    table2 = doc.add_table(rows=1, cols=4)
    table2.style = 'Light Grid Accent 1'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers2 = ['Этап', 'Описание', 'Return', 'Эффект']
    for i, header in enumerate(headers2):
        table2.rows[0].cells[i].text = header
        for paragraph in table2.rows[0].cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    for stage, desc, ret, effect in stages:
        row_cells = table2.add_row().cells
        row_cells[0].text = stage
        row_cells[1].text = desc
        row_cells[2].text = ret
        row_cells[3].text = effect

    doc.add_page_break()

    # Walk-Forward Results
    doc.add_heading('Walk-Forward результаты (TCN)', level=1)

    wf_results = [
        ('1', '+9.49%', '-3.14%', '35.3%', '1.59'),
        ('2', '-0.26%', '-0.26%', '0.0%', '0.00'),
        ('3', '+0.72%', '-0.65%', '42.9%', '1.95'),
        ('4', '-0.63%', '-1.83%', '27.3%', '0.91'),
        ('5', '+1.89%', '-4.08%', '36.0%', '1.20'),
        ('6', '+13.70%', '-3.45%', '42.9%', '2.09'),
        ('Итого', '+26.62%', '-4.08%', '30.7%', '-'),
    ]

    table3 = doc.add_table(rows=1, cols=5)
    table3.style = 'Light Grid Accent 1'
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers3 = ['Fold', 'Return', 'Drawdown', 'Win Rate', 'PF']
    for i, header in enumerate(headers3):
        table3.rows[0].cells[i].text = header
        for paragraph in table3.rows[0].cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    for fold, ret, dd, wr, pf in wf_results:
        row_cells = table3.add_row().cells
        row_cells[0].text = fold
        row_cells[1].text = ret
        row_cells[2].text = dd
        row_cells[3].text = wr
        row_cells[4].text = pf
        if fold == 'Итого':
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

    doc.add_page_break()

    # Key Improvements
    doc.add_heading('Ключевые улучшения', level=1)

    doc.add_heading('1. TA-Lib индикаторы (60+)', level=2)
    indicators = [
        'Trend: EMA, MACD, ADX, Aroon, SAR',
        'Momentum: RSI, Stochastic, Williams %R, CCI',
        'Volatility: Bollinger Bands, ATR, Keltner Channel',
        'Volume: OBV, AD, MFI',
        'Patterns: Doji, Hammer, Engulfing',
    ]
    for ind in indicators:
        doc.add_paragraph(ind, style='List Bullet')

    doc.add_heading('2. Trailing Stop (1%)', level=2)
    doc.add_paragraph('Фиксирует прибыль при движении цены вверх')
    doc.add_paragraph('Снижает drawdown на 70%')

    doc.add_heading('3. Dynamic Position Sizing', level=2)
    doc.add_paragraph('Уменьшает размер позиции при высокой волатильности')
    doc.add_paragraph('Использует ATR для расчета')

    doc.add_heading('4. TCN архитектура', level=2)
    doc.add_paragraph('Temporal Convolutional Network')
    doc.add_paragraph('Profit Factor: 1.72')

    doc.add_page_break()

    # Best Parameters
    doc.add_heading('Лучшие параметры', level=1)

    params = [
        ('model_type', 'tcn'),
        ('hidden_size', '32'),
        ('num_layers', '1'),
        ('dropout', '0.395'),
        ('window', '30'),
        ('batch_size', '64'),
        ('learning_rate', '0.000205'),
        ('n_models', '3'),
        ('stop_loss_pct', '3.07%'),
        ('take_profit_pct', '4.54%'),
        ('trailing_stop', '1%'),
        ('dynamic_sizing', 'ON'),
    ]

    table4 = doc.add_table(rows=1, cols=2)
    table4.style = 'Light Grid Accent 1'
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER

    table4.rows[0].cells[0].text = 'Параметр'
    table4.rows[0].cells[1].text = 'Значение'
    for cell in table4.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    for param, value in params:
        row_cells = table4.add_row().cells
        row_cells[0].text = param
        row_cells[1].text = value

    doc.add_page_break()

    # Conclusions
    doc.add_heading('Выводы', level=1)

    conclusions = [
        'TCN модель показала лучшие результаты (+26.62%)',
        'Trailing Stop снизил drawdown на 70%',
        'Dynamic Sizing улучшил стабильность',
        'TA-Lib дал 60+ индикаторов',
        'Walk-forward подтвердил обобщающую способность',
    ]

    for conclusion in conclusions:
        doc.add_paragraph(conclusion, style='List Number')

    # Save
    output_path = os.path.join(os.path.dirname(__file__), 'Neural_Trading_Final_Report.docx')
    doc.save(output_path)
    print(f'Final report saved to: {output_path}')
    return output_path


if __name__ == '__main__':
    create_final_report()
