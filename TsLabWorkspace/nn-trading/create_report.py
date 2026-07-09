"""Create NN-Trading project report in docx format."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path


def create_report():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # Title
    title = doc.add_heading('Нейросетевая торговая стратегия', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Ensemble TCN + LSTM + Transformer с оптимальным подбором весов')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    doc.add_paragraph()

    # 1. Описание проекта
    doc.add_heading('1. Описание проекта', level=1)
    doc.add_paragraph(
        'Проект представляет собой нейросетевую торговую стратегию для криптовалют, '
        'использующую ансамбль из трёх архитектур: TCN (Temporal Convolutional Network), '
        'LSTM (Long Short-Term Memory) и Transformer. Стратегия включает оптимальный '
        'подбор весов моделей на валидационной выборке, адаптивный порог сигналов и '
        'продвинутый риск-менеджмент с trailing stop и динамическим размером позиции.'
    )

    # 2. Архитектура
    doc.add_heading('2. Архитектура системы', level=1)

    doc.add_heading('2.1 Базовые модели', level=2)

    table = doc.add_table(rows=4, cols=3)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['Модель', 'Описание', 'Роль в ансамбле']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    data = [
        ['TCN', 'Temporal Convolutional Network — свёрточная сеть с дилатированными свёртками', 'Лучше для коротких паттернов'],
        ['LSTM', 'Long Short-Term Memory — рекуррентная сеть с вентилями', 'Лучше для долгосрочных зависимостей'],
        ['Transformer', 'Self-attention механизм — параллельная обработка последовательностей', 'Лучше для глобальных паттернов'],
    ]
    for i, row in enumerate(data):
        for j, val in enumerate(row):
            table.rows[i+1].cells[j].text = val

    doc.add_paragraph()

    doc.add_heading('2.2 Ансамбль с оптимальными весами', level=2)
    doc.add_paragraph(
        'Веса моделей подбираются на валидационной выборке методом grid search '
        'по комбинациям [0.1, 0.2, 0.3, 0.4, 0.5] с условием sum=1. '
        'Для каждой комбинации считается доходность на val, выбираются веса '
        'с максимальным результатом.'
    )

    doc.add_heading('2.3 Оптимальный порог', level=2)
    doc.add_paragraph(
        'Порог вероятности для генерации сигналов подбирается на валидации '
        'из набора [0.3, 0.35, 0.4, 0.45, 0.5]. Это позволяет адаптироваться '
        'к разным рыночным условиям на каждом фолде.'
    )

    # 3. Признаки
    doc.add_heading('3. Технические индикаторы (60+)', level=1)

    doc.add_heading('3.1 Trend', level=2)
    doc.add_paragraph('EMA (10, 20, 50), MACD, ADX, Aroon, Parabolic SAR')

    doc.add_heading('3.2 Momentum', level=2)
    doc.add_paragraph('RSI, Stochastic, Williams %R, CCI, ROC, Momentum, Ultimate Oscillator')

    doc.add_heading('3.3 Volatility', level=2)
    doc.add_paragraph('Bollinger Bands, ATR, Keltner Channel, NATR')

    doc.add_heading('3.4 Volume', level=2)
    doc.add_paragraph('OBV, AD, MFI, Volume SMA, Volume Ratio')

    doc.add_heading('3.5 Price Patterns', level=2)
    doc.add_paragraph('Doji, Hammer, Engulfing, Morning Star, Evening Star')

    # 4. Результаты
    doc.add_heading('4. Результаты Walk-Forward валидации', level=1)

    table2 = doc.add_table(rows=8, cols=6)
    table2.style = 'Light Grid Accent 1'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers2 = ['Фолд', 'TCN вес', 'LSTM вес', 'Transformer вес', 'Return', 'Drawdown']
    for i, h in enumerate(headers2):
        table2.rows[0].cells[i].text = h

    results = [
        ['1', '0.4', '0.1', '0.5', '+2.61%', '-2.96%'],
        ['2', '0.3', '0.3', '0.3', '+10.65%', '-8.08%'],
        ['3', '0.3', '0.3', '0.3', '+21.93%', '-4.35%'],
        ['4', '0.2', '0.5', '0.3', '+3.74%', '-4.01%'],
        ['5', '0.4', '0.2', '0.4', '-9.54%', '-9.58%'],
        ['6', '0.2', '0.5', '0.3', '+2.42%', '-1.38%'],
        ['ИТОГО', '—', '—', '—', '+33.07%', '-13.10%'],
    ]
    for i, row in enumerate(results):
        for j, val in enumerate(row):
            table2.rows[i+1].cells[j].text = val
            if i == 6:  # итого строка
                for run in table2.rows[i+1].cells[j].paragraphs[0].runs:
                    run.bold = True

    doc.add_paragraph()

    # 5. Сравнение подходов
    doc.add_heading('5. Сравнение подходов', level=1)

    table3 = doc.add_table(rows=4, cols=4)
    table3.style = 'Light Grid Accent 1'
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers3 = ['Подход', 'Total Return', 'Max DD', 'Win Rate']
    for i, h in enumerate(headers3):
        table3.rows[0].cells[i].text = h

    approaches = [
        ['Weighted Averaging (val_acc)', '+27.53%', '-15.40%', '36.2%'],
        ['Stacking (MLP meta)', '+16.81%', '-21.09%', '35.6%'],
        ['Optimal Weight Search', '+33.07%', '-13.10%', '38.1%'],
    ]
    for i, row in enumerate(approaches):
        for j, val in enumerate(row):
            table3.rows[i+1].cells[j].text = val

    doc.add_paragraph()

    # 6. Риск-менеджмент
    doc.add_heading('6. Риск-менеджмент', level=1)

    rm_data = [
        ('Stop Loss', '3.07%'),
        ('Take Profit', '4.54%'),
        ('Trailing Stop', '1%'),
        ('Dynamic Sizing', 'Включено (на основе ATR)'),
        ('Risk per Trade', '2% от equity'),
    ]
    for label, value in rm_data:
        p = doc.add_paragraph()
        run = p.add_run(f'{label}: ')
        run.bold = True
        p.add_run(value)

    # 7. Лучшие параметры
    doc.add_heading('7. Оптимальные параметры', level=1)

    params = [
        ('Model Type', 'TCN / LSTM / Transformer (ансамбль)'),
        ('Hidden Size', '32'),
        ('Num Layers', '1'),
        ('Dropout', '0.395'),
        ('Window', '30'),
        ('Batch Size', '64'),
        ('Learning Rate', '0.000205'),
        ('N Models', '3'),
        ('Ensemble Method', 'Optimal Weight Search'),
    ]
    for label, value in params:
        p = doc.add_paragraph()
        run = p.add_run(f'{label}: ')
        run.bold = True
        p.add_run(value)

    # 8. Заключение
    doc.add_heading('8. Заключение', level=1)
    doc.add_paragraph(
        'Нейросетевая стратегия с ансамблем TCN+LSTM+Transformer и оптимальным '
        'подбором весов показала доходность +33.07% за 6 фолдов walk-forward '
        'валидации при максимальном просадке -13.10%. Оптимальный подбор весов '
        'обеспечил улучшение на +5.54% по доходности и -2.3% по drawdown '
        'по сравнению с простым взвешиванием по accuracy.'
    )

    doc.add_paragraph(
        'Ключевые факторы успеха: '
        '(1) разнообразие архитектур в ансамбле, '
        '(2) подбор весов на валидации, '
        '(3) адаптивный порог сигналов, '
        '(4) trailing stop + dynamic sizing для риск-менеджмента.'
    )

    # Save
    output_path = Path("C:/Users/i59400f/Desktop/ai-agent/TsLabWorkspace/nn-trading/Neural_Trading_Report.docx")
    doc.save(str(output_path))
    print(f"Report saved to {output_path}")


if __name__ == "__main__":
    create_report()
