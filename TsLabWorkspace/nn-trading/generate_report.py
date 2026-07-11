"""Generate detailed report in docx format (Russian)."""
import json
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call(["pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    return h


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)
    return table


def main():
    doc = Document()

    # Title
    title = doc.add_heading("Neural Trading System v2", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph("Отчёт по результатам разработки и тестирования\nнейросетевой торговой системы на основе данных Binance")
    doc.add_paragraph("Дата: Июль 2026")

    # 1. Summary
    add_heading(doc, "1. Резюме")
    doc.add_paragraph(
        "Разработана гибридная торговая система, совмещающая нейросетевые модели "
        "(LSTM, Transformer, Attention) с классическими ML (XGBoost, CatBoost, LightGBM, RandomForest). "
        "Система протестирована на 11 инструментах cryptocurrency с использованием "
        "мультитаймфреймового анализа (1d + 1w + 1M) и walk-forward валидации."
    )

    # 2. Data
    add_heading(doc, "2. Данные")
    doc.add_paragraph("Источник: Binance mainnet API (публичный)")
    doc.add_paragraph("Инструменты: BTC, ETH, SOL, NEAR, XLM, AAVE, LINK, SUI, ADA, BCH, TRX, UNI")
    doc.add_paragraph("Таймфреймы: 1d (1820 баров), 1w (260), 1M (60)")
    doc.add_paragraph("Для BTC дополнительно: 5m, 15m, 30m, 1h, 4h")

    # 3. Architecture
    add_heading(doc, "3. Архитектура моделей")

    add_heading(doc, "3.1 HybridEnsemble", level=2)
    doc.add_paragraph(
        "Ансамбль из 7 моделей:\n"
        "- 3 нейросети (LSTM/Transformer/Attention) с разными seed\n"
        "- XGBoost, CatBoost, LightGBM, RandomForest\n"
        "Веса: оптимизированы через grid search на validation"
    )

    add_heading(doc, "3.2 Мультитаймфреймовые фичи", level=2)
    doc.add_paragraph(
        "Для каждого таймфрейма считаются:\n"
        "EMA(10,20,50), RSI(14), MACD, Bollinger Bands, ATR, ADX,\n"
        "Stochastic, Williams %R, CCI, ROC, Momentum, OBV, MFI\n"
        "Итого: 50-170 фичей → auto-select топ-50"
    )

    add_heading(doc, "3.3 Walk-forward валидация", level=2)
    doc.add_paragraph(
        "Anchored (расширяющееся) окно:\n"
        "- Train: от начала до fold_size*(i+1)\n"
        "- Test: следующие fold_size баров\n"
        "- 5 фолдов для каждого инструмента"
    )

    # 4. Results
    add_heading(doc, "4. Результаты")

    # Load 10-instrument results
    results_path = Path("results/instruments_walkforward.json")
    if results_path.exists():
        with open(results_path) as f:
            data = json.load(f)
        instruments = data.get("instruments", [])

        add_heading(doc, "4.1 Walk-forward: 10 инструментов (3 фолда)", level=2)
        headers = ["Инструмент", "Return", "Drawdown", "Win Rate", "Sharpe", "Sortino"]
        rows = []
        for inst in instruments:
            rows.append([
                inst["instrument"],
                f"{inst['total_return']:+.2f}%",
                f"{inst['max_drawdown']:.2f}%",
                f"{inst['avg_win_rate']:.1f}%",
                f"{inst['sharpe']:.2f}",
                f"{inst['sortino']:.2f}",
            ])
        add_table(doc, headers, rows)

        profitable = [i for i in instruments if i["total_return"] > 0]
        avg_ret = sum(i["total_return"] for i in profitable) / len(profitable) if profitable else 0
        avg_sharpe = sum(i["sharpe"] for i in profitable) / len(profitable) if profitable else 0

        doc.add_paragraph(f"\nПрибыльных инструментов: {len(profitable)}/{len(instruments)}")
        doc.add_paragraph(f"Средняя доходность (прибыльных): {avg_ret:+.2f}%")
        doc.add_paragraph(f"Средний Sharpe (прибыльных): {avg_sharpe:+.2f}")

    # Load 3-year backtest results
    bt_path = Path("results/full_backtest_3yr.json")
    if bt_path.exists():
        with open(bt_path) as f:
            bt_data = json.load(f)
        portfolio = bt_data.get("portfolio", {})

        add_heading(doc, "4.2 3-летний бэктест портфеля", level=2)
        doc.add_paragraph(f"Total Return: {portfolio.get('total_return', 0):+.2f}%")
        doc.add_paragraph(f"Max Drawdown: {portfolio.get('max_drawdown', 0):.2f}%")
        doc.add_paragraph(f"Sharpe Ratio: {portfolio.get('sharpe', 0):.2f}")
        doc.add_paragraph(f"Sortino Ratio: {portfolio.get('sortino', 0):.2f}")

    # 5. Portfolio
    add_heading(doc, "5. Portfolio Analysis")
    alloc_path = Path("results/portfolio_allocation.json")
    if alloc_path.exists():
        with open(alloc_path) as f:
            alloc_data = json.load(f)
        alloc = alloc_data.get("allocation", {})
        doc.add_paragraph(f"Ожидаемая доходность портфеля: {alloc_data.get('expected_return', 0):+.2f}%")
        doc.add_paragraph("Распределение:")
        for name, pct in sorted(alloc.items(), key=lambda x: -x[1]):
            doc.add_paragraph(f"  {name}: {pct}%", style="List Bullet")

    # 6. Recommendations
    add_heading(doc, "6. Рекомендации")
    doc.add_paragraph(
        "1. XLM, NEAR, SOL, AAVE, BCH — прибыльные инструменты, основа портфеля\n"
        "2. ADA, LINK, UNI, TRX — убыточны, исключить из портфеля\n"
        "3. Динамическая ребалансировка каждые 30 баров\n"
        "4. Confidence threshold 0.5+ для снижения количества сделок\n"
        "5. Мониторинг drawdown — при превышении 15% снижать позиции\n"
        "6. BTC требует further оптимизации (RL, больше данных)"
    )

    # 7. Technical details
    add_heading(doc, "7. Технические детали")
    doc.add_paragraph(
        "- Python 3.13, PyTorch 2.x, scikit-learn, XGBoost, CatBoost, LightGBM\n"
        "- CPU throttling: 2 потока (OMP_NUM_THREADS=2)\n"
        "- Комиссии: 0.1% taker + 0.05% slippage + 0.05% spread\n"
        "- Risk management: dynamic position sizing based on ATR"
    )

    out_path = Path("NN_Trading_Report_v2.docx")
    doc.save(str(out_path))
    print(f"Report saved to {out_path}")


if __name__ == "__main__":
    main()
