from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# Стили
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Заголовок
title = doc.add_heading('NN Trading Project', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')
subtitle = doc.add_paragraph('C# инференс + TorchSharp + Binance Testnet Trading')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# Содержание
doc.add_heading('Содержание', 1)
doc.add_paragraph('1. Обзор проекта')
doc.add_paragraph('2. Архитектура')
doc.add_paragraph('3. TorchSharp инференс')
doc.add_paragraph('4. Торговые сигналы')
doc.add_paragraph('5. Binance Testnet Trading')
doc.add_paragraph('6. Результаты')
doc.add_paragraph('7. Использование')

doc.add_page_break()

# 1. Обзор проекта
doc.add_heading('1. Обзор проекта', 1)
doc.add_paragraph(
    'NN Trading — это система для торговли криптовалютами на основе нейронных сетей. '
    'Проект включает обучение моделей на Python, экспорт в формат TorchSharp/ONNX, '
    'инференс на C# и исполнение ордеров через Binance Testnet API.'
)

# 2. Архитектура
doc.add_heading('2. Архитектура', 1)
doc.add_paragraph(
    'Система состоит из трёх основных компонентов:'
)

# Таблица архитектуры
table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Компонент', 'Технология', 'Описание']
for i, header in enumerate(headers):
    table.rows[0].cells[i].text = header
    table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

data = [
    ['Обучение', 'Python + PyTorch', 'LSTM модель, 50+ индикаторов, walk-forward'],
    ['Инференс', 'C# + TorchSharp', 'Загрузка .pth моделей, индикаторы на C#'],
    ['Торговля', 'C# + HttpClient', 'Binance REST API, HMAC-SHA256 подпись'],
]
for i, row_data in enumerate(data):
    for j, cell_data in enumerate(row_data):
        table.rows[i+1].cells[j].text = cell_data

doc.add_paragraph('')

# 3. TorchSharp инференс
doc.add_heading('3. TorchSharp инференс', 1)
doc.add_paragraph(
    'TorchSharp позволяет загружать PyTorch модели (.pth) напрямую в C# без экспорта в ONNX. '
    'Конвертация .pth → .npy через Python скрипт export_torchsharp.py.'
)

doc.add_heading('Компоненты:', 2)
doc.add_paragraph('• TorchSharpEngine — загрузка .npy весов, LSTM инференс', style='List Bullet')
doc.add_paragraph('• OnnxEngine — fallback через ONNX Runtime', style='List Bullet')
doc.add_paragraph('• FeatureComputer — 10 индикаторов на C#', style='List Bullet')
doc.add_paragraph('• StandardScaler — Z-нормализация', style='List Bullet')

# 4. Торговые сигналы
doc.add_heading('4. Торговые сигналы', 1)
doc.add_paragraph(
    'Модуль TradingSignals генерирует сигналы на основе предсказаний модели:'
)

doc.add_heading('Входы:', 2)
doc.add_paragraph('• Prediction — предсказание модели (0-1)', style='List Bullet')
doc.add_paragraph('• OHLCV данные — свечи с Binance', style='List Bullet')
doc.add_paragraph('• Текущая позиция — Long/Short/None', style='List Bullet')

doc.add_heading('Выходы:', 2)
doc.add_paragraph('• SignalType — Long/Short/CloseLong/CloseShort/None', style='List Bullet')
doc.add_paragraph('• SignalStrength — 1-5 (Weak → VeryStrong)', style='List Bullet')
doc.add_paragraph('• Stop Loss — ATR-based', style='List Bullet')
doc.add_paragraph('• Take Profit — ATR-based', style='List Bullet')
doc.add_paragraph('• Position Size — Kelly criterion', style='List Bullet')

# 5. Binance Testnet Trading
doc.add_heading('5. Binance Testnet Trading', 1)
doc.add_paragraph(
    'Торговля через Binance Testnet API (https://testnet.binance.vision/). '
    'Поддержка рыночных ордеров, Stop Loss, Take Profit, OCO ордеров.'
)

doc.add_heading('Команды:', 2)
doc.add_paragraph('• dotnet run -- --single — один символ', style='List Bullet')
doc.add_paragraph('• dotnet run -- --scan — мульти-символьный сканер', style='List Bullet')
doc.add_paragraph('• dotnet run -- --trade — торговля с подтверждением', style='List Bullet')

# 6. Результаты
doc.add_heading('6. Результаты', 1)

table2 = doc.add_table(rows=4, cols=3)
table2.style = 'Light Grid Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

headers2 = ['Сделка', 'Цена', 'P&L']
for i, header in enumerate(headers2):
    table2.rows[0].cells[i].text = header
    table2.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

trades = [
    ['BUY 0.00309 BTC', '$64,550', '+$199.46'],
    ['SELL 1.00309 BTC', '$64,514', '+$64,713.34'],
    ['BUY 20,030 DOGE', '$0.0744', '+$1,877.87'],
]
for i, row_data in enumerate(trades):
    for j, cell_data in enumerate(row_data):
        table2.rows[i+1].cells[j].text = cell_data

doc.add_paragraph('')
doc.add_paragraph('Итоговый баланс: $74,901.51 USDT + 1 ETH')

# 7. Использование
doc.add_heading('7. Использование', 1)

doc.add_heading('Быстрый старт:', 2)
doc.add_paragraph('1. Установить .env файл с API ключами', style='List Number')
doc.add_paragraph('2. Запустить dotnet run -- --trade --symbols BTCUSDT', style='List Number')
doc.add_paragraph('3. Подтвердить сделку (y/n)', style='List Number')

doc.add_heading('Python скрипты:', 2)
doc.add_paragraph('• scripts/check_status.py — проверка баланса и ордеров', style='List Bullet')
doc.add_paragraph('• scripts/buy_doge.py — покупка DOGE', style='List Bullet')
doc.add_paragraph('• scripts/sell_btc.py — продажа BTC', style='List Bullet')
doc.add_paragraph('• scripts/set_sl_tp.py — установка SL/TP', style='List Bullet')

# Сохранение
output_path = os.path.join(os.path.dirname(__file__), '..', 'NN_Trading_Description.docx')
doc.save(output_path)
print(f'Документ сохранён: {output_path}')
