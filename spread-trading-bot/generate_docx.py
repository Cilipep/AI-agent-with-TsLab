"""Generate bot description and summary docx in Russian"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime


def setup_page(doc):
    section = doc.sections[0]
    section.page_width, section.page_height = Cm(21.0), Cm(29.7)
    section.top_margin = section.bottom_margin = Cm(2.54)
    section.left_margin = section.right_margin = Cm(3.18)
    section.orientation = WD_ORIENTATION.PORTRAIT


def tune_styles(doc):
    body = doc.styles["Normal"]
    body.font.name = "Calibri"
    body.font.size = Pt(11)
    body.paragraph_format.line_spacing = 1.15
    body.paragraph_format.space_after = Pt(6)

    for n, size in [(1, 18), (2, 14), (3, 12)]:
        s = doc.styles[f"Heading {n}"]
        s.font.name = "Calibri Light"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
        s.paragraph_format.space_before = Pt(14 - 2 * n)
        s.paragraph_format.space_after = Pt(4)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)

    for row_idx, row_data in enumerate(rows):
        for col_idx, val in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    return table


def main():
    doc = Document()
    setup_page(doc)
    tune_styles(doc)

    # Title
    title = doc.add_paragraph("Spread Trading Bot", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("Bybit BTCUSDT/ETHUSDT", style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_p = doc.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y')}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # 1. Description
    doc.add_heading("1. Описание бота", level=1)

    doc.add_paragraph(
        "Spread Trading Bot — автоматическая торговая система для биржи Bybit, "
        "реализующая стредж-трейдинг на паре BTCUSDT/ETHUSDT. Бот основан на "
        "принципах Александра Герчика: строгие стоп-лоссы, торговый журнал, "
        "ограничение дневных убытков и адаптация к волатильности."
    )

    doc.add_heading("Принцип работы", level=2)
    doc.add_paragraph(
        "Бот рассчитывает спред (отношение цен) между BTC и ETH, вычисляет "
        "скользящее среднее и стандартное отклонение, затем определяет Z-Score — "
        "отклонение текущего спреда от нормы. Когда Z-Score превышает порог входа, "
        "бот открывает позицию (longspread или shortspread), ожидая возврата "
        "спреда к среднему."
    )

    doc.add_heading("Управление рисками", level=2)
    items = [
        "Стоп-лосс по Z-Score (автоматический выход при отклонении)",
        "Ограничение дневных убытков (2% от баланса)",
        "Ограничение количества сделок в день (10)",
        "Максимальная просадка (10%)",
        "Серверные TP/SL ордера (работают даже при выключенном боте)",
        "Адаптация размера позиции к волатильности",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Технические характеристики", level=2)
    add_table(doc, ["Параметр", "Значение"], [
        ["Торговая площадка", "Bybit (mainnet)"],
        ["Пара инструментов", "BTCUSDT / ETHUSDT"],
        ["Таймфрейм", "4 часа (оптимизирован)"],
        ["ENTRY_Z", "1.2"],
        ["EXIT_Z", "0.5"],
        ["STOP_Z", "3.5"],
        ["Размер позиции", "5 USDT за ногу"],
        ["Плечо", "10x"],
        ["Lookback", "100 свечей"],
    ])

    doc.add_page_break()

    # 2. Optimization
    doc.add_heading("2. Оптимизация параметров", level=1)

    doc.add_paragraph(
        "Параметры ENTRY_Z и STOP_Z были оптимизированы методом сеточного "
        "перебора (grid search) на исторических данных за ~167 дней (1000 свечей 4h)."
    )

    doc.add_heading("Результаты оптимизации по таймфреймам", level=2)
    add_table(doc, ["TF", "ENTRY_Z", "STOP_Z", "PnL%", "Trades", "Win Rate", "Max DD", "PF"], [
        ["4h", "1.2", "3.5", "+3.6%", "12", "58.3%", "1.1%", "1.76"],
        ["1h", "2.5", "5.0", "+1.6%", "4", "100%", "0.0%", "99.0"],
        ["1d", "2.5", "5.0", "+1.5%", "2", "100%", "0.0%", "99.0"],
        ["15m", "2.5", "3.5", "+0.7%", "4", "100%", "0.0%", "99.0"],
        ["1m", "2.5", "4.0", "+0.1%", "11", "63.6%", "0.2%", "1.20"],
        ["5m", "0.8", "2.0", "+0.1%", "105", "49.5%", "0.6%", "1.02"],
    ])

    doc.add_paragraph(
        "4h выбран как лучший таймфрейм: максимальная прибыль (+3.6%), "
        "достаточное количество сделок (12), приемлемый риск (Max DD 1.1%)."
    )

    doc.add_page_break()

    # 3. Full backtest
    doc.add_heading("3. Полный бэктест на 4h", level=1)

    doc.add_paragraph(
        "Полный бэктест на 1000 свечей 4h (период: 06.02.2026 — 22.07.2026, "
        "~167 дней) с оптимальными параметрами ENTRY_Z=1.2, EXIT_Z=0.5, STOP_Z=3.5."
    )

    doc.add_heading("Результаты", level=2)
    add_table(doc, ["Метрика", "Значение"], [
        ["Начальный баланс", "$100.00"],
        ["Конечный баланс", "$105.23"],
        ["Total PnL", "+$5.23 (+5.2%)"],
        ["Max Drawdown", "5.1%"],
        ["Total Trades", "26"],
        ["Win Rate", "69.2%"],
        ["Profit Factor", "1.44"],
        ["Avg PnL/Trade", "$0.20"],
    ])

    doc.add_heading("Анализ", level=2)
    doc.add_paragraph(
        "Стратегия показывает стабильную прибыльность на исторических данных. "
        "Win Rate 69.2% и Profit Factor 1.44 указывают на положительное "
        "ожидание. Максимальная просадка 5.1% находится в пределах допустимого "
        "риска (10% MAX_DRAWDOWN). Средняя прибыль на сделку $0.20 при размере "
        "позиции 5 USDT x 10x = 50 USDT."
    )

    doc.add_page_break()

    # 4. Files
    doc.add_heading("4. Файлы проекта", level=1)
    add_table(doc, ["Файл", "Описание"], [
        ["spotread_bot.py", "Основной файл бота (торговая логика)"],
        ["config.py", "Конфигурация (API ключи, параметры, лимиты)"],
        ["backtest.py", "Скрипт бэктеста и оптимизации параметров"],
        ["full_backtest_4h.py", "Полный бэктест на 4h"],
        ["test_bot.py", "Тестирование API соединения"],
        ["spread_bot.log", "Лог работы бота"],
        ["trade_journal.csv", "Журнал сделок"],
        ["daily_report.txt", "Дневные отчёты"],
    ])

    doc.add_page_break()

    # 5. Recommendations
    doc.add_heading("5. Рекомендации", level=1)

    items = [
        "Запускать бота на 4h таймфрейме с параметрами ENTRY_Z=1.2, STOP_Z=3.5",
        "Мониторить журнал сделок (trade_journal.csv) и логи (spread_bot.log)",
        "Регулярно проводить повторную оптимизацию на свежих данных",
        "Рассмотреть увеличение размера позиции при стабильной прибыльности",
        "Не превышать 10% максимальной просадки — при достижении остановить бота",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Number")

    doc.save("Spread_Trading_Bot_Description.docx")
    print("Saved: Spread_Trading_Bot_Description.docx")


if __name__ == "__main__":
    main()
