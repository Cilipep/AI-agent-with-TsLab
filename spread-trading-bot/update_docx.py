"""Add multi-instrument test results to existing docx"""
from docx import Document
from docx.shared import Pt, RGBColor


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)

    for row_idx, row_data in enumerate(rows):
        for col_idx, val in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    return table


def main():
    doc = Document("Spread_Trading_Bot_Description.docx")

    # Find section 4 (files) to insert before recommendations
    insert_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.startswith("5. Рекомендации"):
            insert_idx = i
            break

    if insert_idx is None:
        print("Could not find recommendations section")
        return

    # We need to add content before section 5
    # Since python-docx doesn't support inserting at index easily,
    # we'll regenerate the full document with the new section

    # Read all existing content
    paragraphs = []
    for p in doc.paragraphs:
        paragraphs.append({
            "text": p.text,
            "style": p.style.name
        })

    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            rows.append(cells)
        tables.append(rows)

    # Find where section 5 starts in paragraphs
    sec5_idx = None
    for i, p in enumerate(paragraphs):
        if p["text"].startswith("5. Рекомендации"):
            sec5_idx = i
            break

    # Rebuild document
    doc2 = Document()

    # Copy styles from original
    section = doc2.sections[0]
    from docx.shared import Cm
    section.page_width, section.page_height = Cm(21.0), Cm(29.7)
    section.top_margin = section.bottom_margin = Cm(2.54)
    section.left_margin = section.right_margin = Cm(3.18)

    # Add paragraphs before section 5
    table_counter = 0
    for i in range(sec5_idx):
        p_info = paragraphs[i]
        if p_info["style"] == "Normal" and p_info["text"] == "":
            # Check if next paragraph has a table
            if i + 1 < sec5_idx and paragraphs[i + 1]["style"] == "Normal":
                # This might be before a table
                pass
            doc2.add_paragraph("")
        elif "Heading" in p_info["style"]:
            doc2.add_paragraph(p_info["text"], style=p_info["style"])
        elif p_info["style"] == "List Bullet":
            doc2.add_paragraph(p_info["text"], style="List Bullet")
        elif p_info["style"] == "List Number":
            doc2.add_paragraph(p_info["text"], style="List Number")
        elif p_info["text"]:
            doc2.add_paragraph(p_info["text"], style="Normal")

    # Add new section: Multi-Instrument Test
    doc2.add_page_break()
    doc2.add_heading("6. Тестирование на других инструментах", level=1)

    doc2.add_paragraph(
        "Стратегия была протестирована на различных торговых парах биржи Bybit "
        "с оптимальными параметрами (ENTRY_Z=1.2, EXIT_Z=0.5, STOP_Z=3.5, 4h)."
    )

    doc2.add_heading("Результаты", level=2)

    instrument_data = [
        ["BTCUSDT/ETHUSDT", "+5.2%", "26", "69.2%", "5.1%", "1.44", "Лучший"],
        ["ETHUSDT/BTCUSDT", "+5.0%", "26", "65.4%", "4.6%", "1.42", "Хороший"],
        ["BTCUSDT/SOLUSDT", "+0.6%", "22", "59.1%", "5.4%", "1.06", "Пограничный"],
        ["SOLUSDT/BTCUSDT", "+0.4%", "22", "54.5%", "6.4%", "1.03", "Пограничный"],
        ["DOGEUSDT/BTCUSDT", "-0.0%", "21", "76.2%", "3.3%", "1.00", "Нулевой"],
        ["BTCUSDT/DOGEUSDT", "-0.9%", "20", "75.0%", "3.5%", "0.93", "Убыточный"],
        ["ETHUSDT/SOLUSDT", "-4.1%", "26", "73.1%", "7.7%", "0.77", "Убыточный"],
        ["SOLUSDT/ETHUSDT", "-4.2%", "27", "70.4%", "8.2%", "0.77", "Убыточный"],
        ["DOGEUSDT/ETHUSDT", "-12.0%", "15", "53.3%", "3.7%", "0.33", "Убыточный"],
        ["ETHUSDT/DOGEUSDT", "-17.3%", "16", "56.2%", "6.8%", "0.30", "Убыточный"],
    ]

    add_table(doc2, ["Pair", "PnL%", "Trades", "Win Rate", "Max DD", "PF", "Статус"], instrument_data)

    doc2.add_heading("Выводы", level=2)

    conclusions = [
        "BTC/ETH — лучшая пара (+5.2%, PF 1.44)",
        "ETH/BTC тоже прибыльна (+5.0%, PF 1.42) — стратегия работает в обе стороны",
        "BTC/SOL и SOL/BTC на грани безубыточности (+0.6%, +0.4%)",
        "Все пары с DOGE убыточные — DOGE слишком волатильный для спред-стратегии",
        "ETH/SOL и SOL/ETH убыточные — недостаточная корреляция",
    ]
    for c in conclusions:
        doc2.add_paragraph(c, style="List Bullet")

    doc2.add_paragraph(
        "Рекомендация: использовать BTCUSDT/ETHUSDT как основную пару. "
        "Можно параллельно запустить ETHUSDT/BTCUSDT как зеркальную стратегию."
    )

    # Add recommendations section (section 5)
    doc2.add_heading("7. Рекомендации", level=1)

    recommendations = [
        "Запускать бота на 4h таймфрейме с параметрами ENTRY_Z=1.2, STOP_Z=3.5",
        "Использовать пару BTCUSDT/ETHUSDT как основную",
        "Мониторить журнал сделок (trade_journal.csv) и логи (spread_bot.log)",
        "Регулярно проводить повторную оптимизацию на свежих данных",
        "Рассмотреть увеличение размера позиции при стабильной прибыльности",
        "Не превышать 10% максимальной просадки — при достижении остановить бота",
    ]
    for i, r in enumerate(recommendations, 1):
        doc2.add_paragraph(r, style="List Number")

    doc2.save("Spread_Trading_Bot_Description.docx")
    print("Updated: Spread_Trading_Bot_Description.docx")


if __name__ == "__main__":
    main()
